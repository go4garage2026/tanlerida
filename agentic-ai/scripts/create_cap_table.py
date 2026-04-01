"""Generate Cap Table & Sole Ownership Statement for Arch Grant"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

today = date.today().strftime("%B %d, %Y")

# Header
h = doc.add_heading('', level=0)
run = h.add_run('CAPITALIZATION TABLE & SOLE OWNERSHIP STATEMENT')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x19, 0x19, 0x19)
run.font.bold = True
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = info.add_run('Tan Lerida (Tangred)\n')
r1.bold = True
r2 = info.add_run('Sole Proprietorship\n')
r2.italic = True
info.add_run('Date: ' + today)

doc.add_paragraph('')

# Section 1: Cap Table
doc.add_heading('1. Capitalization Table', level=1)

table = doc.add_table(rows=4, cols=5)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Shareholder', 'Role', 'Equity Class', 'Ownership %', 'Notes']
for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)

row1 = ['Vivek Pargal', 'Founder & CEO', 'Common', '100%', 'Sole Owner']
for i, val in enumerate(row1):
    table.rows[1].cells[i].text = val

row2 = ['-- Unissued / Reserved', '--', '--', '0%', 'No options pool created']
for i, val in enumerate(row2):
    table.rows[2].cells[i].text = val

row3 = ['TOTAL', '', '', '100%', '']
for i, val in enumerate(row3):
    cell = table.rows[3].cells[i]
    cell.text = val
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

for row in table.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(3.5)
    row.cells[2].width = Cm(2.5)
    row.cells[3].width = Cm(2.5)
    row.cells[4].width = Cm(4)

doc.add_paragraph('')

# Section 2: Ownership Summary
doc.add_heading('2. Ownership Summary', level=1)

bullets = [
    'Total equity holders: 1 (sole proprietor)',
    'Outside investors: None',
    'Outstanding convertible notes: None',
    'SAFE agreements: None',
    'Stock option pool: Not yet established',
    'Debt or liens against the company: None',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph('')

# Section 3: Sole Ownership Statement
doc.add_heading('3. Statement of Sole Ownership', level=1)

statement = doc.add_paragraph()
statement.paragraph_format.space_before = Pt(12)
statement.paragraph_format.space_after = Pt(12)

text = (
    'I, Vivek Pargal, hereby confirm and attest that I am the sole owner '
    'and proprietor of Tan Lerida, operating under the trade name "Tangred." '
    'I hold 100% of the equity interest in the business. There are no other equity holders, '
    'co-founders, investors, or parties with any ownership stake, convertible instruments, '
    'or claims to equity in the company. All intellectual property, assets, and liabilities '
    'of the business are solely held by me as the proprietor. '
    'This statement is true and accurate as of the date signed below.'
)
statement.add_run(text)

doc.add_paragraph('')

facts = [
    ('Legal Entity:', 'Tan Lerida (Sole Proprietorship)'),
    ('Trade Name:', 'Tangred'),
    ('Owner:', 'Vivek Pargal'),
    ('Ownership:', '100% -- Sole Proprietor'),
    ('Outside Equity:', 'None'),
    ('Encumbrances:', 'None'),
]
for label, value in facts:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(2)
    p.paragraph_format.space_after = Pt(2)
    run_l = p.add_run(label + '  ')
    run_l.bold = True
    run_l.font.size = Pt(11)
    run_v = p.add_run(value)
    run_v.font.size = Pt(11)

# Signature block
doc.add_paragraph('')
doc.add_paragraph('')

sig = doc.add_paragraph()
sig.add_run('_' * 45)
sig.paragraph_format.space_after = Pt(2)

name_line = doc.add_paragraph()
r = name_line.add_run('Vivek Pargal')
r.bold = True
name_line.paragraph_format.space_after = Pt(2)

title_line = doc.add_paragraph()
title_line.add_run('Founder & Sole Proprietor, Tan Lerida (d/b/a Tangred)')
title_line.paragraph_format.space_after = Pt(2)

date_line = doc.add_paragraph()
date_line.add_run('Date: ' + today)

# Footer
doc.add_paragraph('')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = footer.add_run('Prepared for: Arch Grant 2026 Startup Competition Application')
run_f.italic = True
run_f.font.size = Pt(9)
run_f.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

out_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'Tangred_Cap_Table_Sole_Ownership_Statement.docx')
doc.save(out_path)
print('Saved:', out_path)
print('Size:', round(os.path.getsize(out_path) / 1024, 1), 'KB')
